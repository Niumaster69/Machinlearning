"""Generate the Word report for Activity 3 — Invisible Maps (K-Means Clustering).
Matches the structure of ML_Activity2_DuvanLozano.pdf.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import Clustering

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_STATIC_DIR = os.path.join(_SCRIPT_DIR, "static")
_OUTPUT = os.path.join(_SCRIPT_DIR, "ML_Activity3_DuvanLozano.docx")

manual = Clustering.manual_result
app_r = Clustering.app_result


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def _shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            bold_prefix, rest = item
            r1 = p.add_run(bold_prefix)
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = p.add_run(rest)
            r2.font.size = Pt(11)
        else:
            run = p.add_run(item)
            run.font.size = Pt(11)


def add_table(doc, headers, rows, header_color="7C3AED", widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        _shade(hdr_cells[i], header_color)

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)

    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)

    return table


def add_spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

doc = Document()

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ============================= COVER PAGE =============================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Unsupervised Machine Learning — K-Means Clustering")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Activity No. 3 — Invisible Maps")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

for _ in range(6):
    doc.add_paragraph()

for line in ["Student", "Duvan Lozano Romero"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(12)
    if line == "Student":
        r.bold = True

for _ in range(6):
    doc.add_paragraph()

for line in [
    "Systems and Computer Engineering Program",
    "University of Cundinamarca — Chia",
    "Course: Machine Learning",
    "April 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(12)

doc.add_page_break()

# ============================ 1. INTRODUCTION ===========================
add_heading(doc, "1. Introduction", 1)

add_heading(doc, "1.1 Description of the Application Update", 2)
add_para(doc,
    "This report documents the third expansion of the Flask web application originally "
    "developed in Activity 1 (Linear Regression) and extended in Activity 2 with Logistic "
    "Regression and K-Nearest Neighbors classification. In this activity, a new Unsupervised "
    "Machine Learning section has been added, dedicated entirely to the K-Means clustering "
    "algorithm."
)
add_para(doc, "The new section includes three sub-pages:")
add_bullets(doc, [
    ("Basic Concepts — ", "Conceptual explanation of unsupervised learning, clustering, and the K-Means algorithm."),
    ("K-Means Manual Exercise — ", "A by-hand simulation of the algorithm on 100 video games across three iterations."),
    ("Clustering Application — ", "A full scikit-learn pipeline on 7,017 video games with interactive visualizations."),
])

add_heading(doc, "1.2 General Objective", 2)
add_para(doc,
    "To analyze and implement the K-Means clustering algorithm through a manual iterative "
    "process and a Flask-based application, demonstrating understanding of unsupervised "
    "learning, distance calculation, centroid updating, cluster assignment, and variance "
    "analysis, while presenting results in a structured and interactive web interface."
)

add_heading(doc, "1.3 Importance of Unsupervised Learning", 2)
add_para(doc,
    "Unsupervised learning is essential whenever the data has no labels. Unlike supervised "
    "algorithms such as Linear Regression, Logistic Regression or KNN, where the model learns "
    "from examples with known answers, unsupervised models discover the inner structure of "
    "data by themselves. Clustering — the most common unsupervised task — groups similar "
    "instances together based on a distance metric."
)
add_para(doc, "In the video game industry, unsupervised learning powers:")
add_bullets(doc, [
    ("Player segmentation — ", "grouping players by behavior for matchmaking and content design."),
    ("Game catalog grouping — ", "discovering natural segments of titles for recommendation systems."),
    ("Anomaly detection — ", "identifying cheaters, bots, or fraudulent purchases."),
    ("Market analysis — ", "profiling commercial versus critical reception to guide release strategies."),
])

doc.add_page_break()

# ============================ 2. APPLICATION STRUCTURE ============================
add_heading(doc, "2. Application Structure", 1)

add_heading(doc, "2.1 Updated Navigation Menu", 2)
add_para(doc, "The menu now contains four main algorithm sections, including the new Unsupervised Machine Learning block:")

menu_rows = [
    ("Home", "—", "/"),
    ("ML Use Cases", "—", "/use-cases"),
    ("Linear Regression", "Basic Concepts", "/linear-regression/concepts"),
    ("", "Application", "/linear-regression/application"),
    ("Logistic Regression", "Basic Concepts", "/logistic-regression/concepts"),
    ("", "Application", "/logistic-regression/application"),
    ("KNN", "Basic Concepts", "/knn/concepts"),
    ("", "Application", "/knn/application"),
    ("Unsupervised Learning", "Basic Concepts", "/unsupervised/concepts"),
    ("", "K-Means Manual Exercise", "/unsupervised/manual"),
    ("", "Clustering Application", "/unsupervised/application"),
]
add_table(doc, ["Menu Item", "Submenu", "Route"], menu_rows, widths=[4.5, 5.5, 6.0])
add_spacer(doc)

add_heading(doc, "2.2 Explanation of the New Section", 2)
add_para(doc, "Unsupervised Learning Section:", bold=True)
add_bullets(doc, [
    ("Basic Concepts — ", "Covers the definition of unsupervised learning, the clustering task, the K-Means algorithm, centroids, Euclidean distance, the iterative process, variance (WCSS), convergence, and real-world applications."),
    ("K-Means Manual Exercise — ", "Presents the Part-1 by-hand simulation on 100 video games with two features, three iterations with distance tables, cluster assignments, centroid updates, and a variance line chart."),
    ("Clustering Application — ", "Presents the K-Means clustering problem on 7,017 video games, data preparation, model training with scikit-learn, cluster summary tables, centroids, a scatter plot with colored clusters, and the interpretation of results."),
])

add_heading(doc, "2.3 Screenshots of the Interface", 2)
add_para(doc, "Home page showing the updated navigation menu with the new Unsupervised Learning dropdown:", bold=True)
add_para(doc, "[ Paste screenshot of the Home page (/) — make sure the Unsupervised Learning dropdown is visible. ]", italic=True)
add_spacer(doc)
add_para(doc, "Unsupervised Learning — Basic Concepts page:", bold=True)
add_para(doc, "[ Paste screenshot of /unsupervised/concepts. ]", italic=True)
add_spacer(doc)
add_para(doc, "K-Means Manual Exercise page:", bold=True)
add_para(doc, "[ Paste screenshot of /unsupervised/manual — ideally showing one of the iteration blocks with the distance table. ]", italic=True)
add_spacer(doc)
add_para(doc, "Clustering Application page:", bold=True)
add_para(doc, "[ Paste screenshot of /unsupervised/application — showing the cluster summary cards or the scatter plot. ]", italic=True)

doc.add_page_break()

# ============================ 3. PART 1 — MANUAL K-MEANS ============================
add_heading(doc, "3. Part 1 — Manual K-Means Simulation", 1)

add_heading(doc, "3.1 Dataset Description", 2)
add_para(doc,
    f"For the manual simulation a random sample of {manual['dataset_size']} games was taken from the "
    "Video Games Sales dataset (Kaggle, 22 Dec 2016). Each record is described by two numerical "
    "features as required by the activity:"
)
add_bullets(doc, [
    ("Critic_Score — ", "Metacritic critical rating on a 0-100 scale."),
    ("Global_Sales — ", "Worldwide sales in millions of units."),
])
add_para(doc, "The goal is to discover K = 3 natural groups of games, based only on these two reception/commercial features, without any external label.")

add_heading(doc, "3.2 Initial Centroids (Manually Chosen)", 2)
add_para(doc, "Three centroids were placed manually, spread across the observed range of the data so that every region has a nearby candidate cluster:")

init_rows = [
    (f"Cluster {i+1}", f"{c['critic_score']}", f"{c['global_sales']}", desc)
    for i, (c, desc) in enumerate(zip(
        manual["initial_centroids"],
        ["Low critic / Low sales", "Mid critic / Mid sales", "High critic / High sales"],
    ))
]
add_table(doc, ["Cluster", "Critic_Score", "Global_Sales", "Description"], init_rows,
          widths=[3.0, 3.5, 3.5, 5.5])
add_spacer(doc)

# Iterations
for it in manual["iterations"]:
    add_heading(doc, f"3.{2 + it['iteration']} Iteration {it['iteration']}", 2)

    add_para(doc, f"Centroids used in Iteration {it['iteration']}:", bold=True)
    iter_in_rows = [
        (f"C{c['cluster']}", f"{c['critic_score']}", f"{c['global_sales']}")
        for c in it["centroids_in"]
    ]
    add_table(doc, ["Cluster", "Critic_Score", "Global_Sales"], iter_in_rows,
              widths=[3.5, 4.5, 4.5])
    add_spacer(doc)

    add_para(doc, "Distance table (sample of first 12 rows — the complete 100-row table is available on the Flask application):", bold=True)
    dist_rows = [
        (r["id"], r["name"][:28], f"{r['critic_score']}", f"{r['global_sales']}",
         f"{r['d1']}", f"{r['d2']}", f"{r['d3']}", f"C{r['cluster']}")
        for r in it["rows"][:12]
    ]
    add_table(doc, ["#", "Game", "Critic", "Sales", "d(C1)", "d(C2)", "d(C3)", "Cluster"],
              dist_rows, widths=[1.0, 5.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0])
    add_spacer(doc)

    add_para(doc, "Cluster counts after assignment:", bold=True)
    add_bullets(doc, [
        f"Cluster 1: {it['counts'][0]} games",
        f"Cluster 2: {it['counts'][1]} games",
        f"Cluster 3: {it['counts'][2]} games",
    ])

    add_para(doc, "New centroids (mean of each cluster):", bold=True)
    new_rows = [
        (f"C{c['cluster']}", f"{c['critic_score']}", f"{c['global_sales']}")
        for c in it["new_centroids"]
    ]
    add_table(doc, ["Cluster", "Critic_Score", "Global_Sales"], new_rows,
              widths=[3.5, 4.5, 4.5])
    add_spacer(doc)

    p = doc.add_paragraph()
    r = p.add_run(f"WCSS (Within-Cluster Sum of Squares) = {it['wcss']}")
    r.bold = True
    r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

    add_spacer(doc)

# Variance comparison
add_heading(doc, "3.6 Variance Comparison", 2)
add_para(doc, "The WCSS decreases monotonically across the three iterations, confirming the algorithm is converging toward a stable solution:")

var_rows = []
for i, it in enumerate(manual["iterations"]):
    change = "—" if i == 0 else f"{it['wcss'] - manual['iterations'][i-1]['wcss']:.3f}"
    var_rows.append((it["iteration"], f"{it['wcss']}", change))
add_table(doc, ["Iteration", "WCSS", "Change"], var_rows, widths=[3.0, 5.0, 5.0])
add_spacer(doc)

variance_png = os.path.join(_STATIC_DIR, "kmeans_variance.png")
if os.path.exists(variance_png):
    doc.add_picture(variance_png, width=Inches(5.5))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run("Figure 1. WCSS reduction across the three manual K-Means iterations.")
    r.italic = True
    r.font.size = Pt(9.5)

add_heading(doc, "3.7 Final Clustering Result", 2)
final = manual["iterations"][-1]
add_para(doc, "After three iterations the clusters identified can be characterized as follows:")
add_bullets(doc, [
    (f"Cluster 1 — ",
     f"centroid at ({final['new_centroids'][0]['critic_score']}, {final['new_centroids'][0]['global_sales']}) — "
     f"{final['counts'][0]} games with low critic scores and modest sales."),
    (f"Cluster 2 — ",
     f"centroid at ({final['new_centroids'][1]['critic_score']}, {final['new_centroids'][1]['global_sales']}) — "
     f"{final['counts'][1]} games with mid critic scores and moderate sales (mainstream bulk)."),
    (f"Cluster 3 — ",
     f"centroid at ({final['new_centroids'][2]['critic_score']}, {final['new_centroids'][2]['global_sales']}) — "
     f"{final['counts'][2]} games with high critic scores and the largest commercial success (blockbusters)."),
])
add_para(doc,
    "The WCSS dropped from {:.3f} to {:.3f} — a reduction of {:.1f}% in just three iterations — "
    "which shows that K-Means makes most of its progress in the very first iterations and then "
    "refines the solution with diminishing returns.".format(
        manual["iterations"][0]["wcss"], final["wcss"],
        100.0 * (manual["iterations"][0]["wcss"] - final["wcss"]) / manual["iterations"][0]["wcss"]
    )
)

doc.add_page_break()

# ============================ 4. CONCEPTS ============================
add_heading(doc, "4. Unsupervised Learning — Concepts", 1)

add_heading(doc, "4.1 Definition of Unsupervised Learning", 2)
add_para(doc,
    "Unsupervised Learning is a branch of machine learning where the algorithm only receives "
    "input data, with no labels or expected outputs. The model must discover the underlying "
    "structure of the data on its own, in contrast to supervised algorithms (Linear Regression, "
    "Logistic Regression, KNN) which learn from input-output pairs."
)

add_heading(doc, "4.2 Clustering", 2)
add_para(doc,
    "Clustering is the most common unsupervised task. Its goal is to partition a dataset into "
    "groups (clusters) such that items inside the same group are more similar to each other "
    "than to items in other groups. Similarity is measured through a distance metric — usually "
    "Euclidean distance."
)

add_heading(doc, "4.3 K-Means Algorithm", 2)
add_para(doc,
    "K-Means is the most widely used clustering algorithm. It partitions a dataset into K "
    "clusters, where K is chosen by the user. Each cluster is represented by a centroid — the "
    "arithmetic mean of all points inside it. The algorithm iteratively refines centroid "
    "positions and point assignments until the clusters stop changing (convergence)."
)

add_heading(doc, "4.4 Centroids and Distance Metric", 2)
add_para(doc, "Centroid: the geometric center of a cluster, equal to the mean of the points it contains. Represented as a vector with one component per feature.")
add_para(doc, "Euclidean distance (the metric K-Means uses):", bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("d(P, C) = √( (x₁ − x₂)² + (y₁ − y₂)² )")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

add_heading(doc, "4.5 Iterative Process", 2)
add_bullets(doc, [
    ("1. Initialize centroids — ", "choose K initial positions (random, manual, or via K-Means++)."),
    ("2. Assign points to nearest centroid — ", "compute Euclidean distance from each point to each centroid and assign it to the closest."),
    ("3. Recalculate centroids — ", "new centroid = mean of the points currently inside its cluster."),
    ("4. Repeat until convergence — ", "loop steps 2-3 until centroids stop moving or the maximum iteration count is reached."),
])

add_heading(doc, "4.6 Variance (WCSS) and Convergence", 2)
add_para(doc,
    "The quantity K-Means minimizes is the Within-Cluster Sum of Squares (WCSS) — the sum of "
    "squared distances from each point to its assigned centroid. WCSS can only decrease (or "
    "stay equal) between iterations. When it stops decreasing, the algorithm has converged."
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("WCSS = Σ Σ d(xᵢ, cₖ)²")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

add_heading(doc, "4.7 Advantages and Limitations", 2)
add_para(doc, "Advantages:", bold=True)
add_bullets(doc, [
    "Simple to understand and implement.",
    "Fast and scales well to large datasets.",
    "Produces interpretable centroids as group prototypes.",
])
add_para(doc, "Limitations:", bold=True)
add_bullets(doc, [
    "Requires K to be set beforehand.",
    "Sensitive to initial centroid placement.",
    "Assumes roughly spherical, similarly-sized clusters.",
    "Requires feature scaling because distances are affected by scale.",
    "May converge to a local minimum rather than the global optimum.",
])

add_heading(doc, "4.8 Real-World Applications", 2)
add_bullets(doc, [
    ("Player segmentation — ", "grouping gamers by behavior for matchmaking and content design."),
    ("Customer segmentation — ", "clustering buyers by purchasing patterns for targeted marketing."),
    ("Anomaly detection — ", "spotting outliers such as fraud, bots, or cheating."),
    ("Image compression — ", "reducing an image's color palette to K representative values."),
    ("Document grouping — ", "topic discovery inside a text corpus."),
])

doc.add_page_break()

# ============================ 5. APPLICATION ============================
add_heading(doc, "5. Clustering Application (Flask + Scikit-learn)", 1)

add_heading(doc, "5.1 Problem Description", 2)
add_para(doc,
    "The video game industry produces two reception signals for every release: the critics' "
    "opinion (Critic_Score) and the players' opinion (User_Score). These two signals often "
    "agree, but sometimes diverge sharply — creating interesting market segments. This "
    "application uses K-Means clustering to discover, without any labels, how games group "
    "themselves when described only by their two reception scores."
)

add_heading(doc, "5.2 Dataset Explanation", 2)
add_bullets(doc, [
    ("Source — ", "Kaggle, Video Games Sales as at 22 Dec 2016."),
    ("Original records — ", "16,719 video games."),
    (f"Records after cleaning — ", f"{app_r['dataset_size']} games with valid Critic_Score, User_Score and Global_Sales (well above the activity's 1,000-record minimum)."),
    ("Features used — ", "2 numerical features (Critic_Score 0-100 and User_Score 0-10)."),
])

add_heading(doc, "5.3 Variables", 2)
add_para(doc, "Input Features:", bold=True)
add_table(doc,
    ["Feature", "Type", "Description"],
    [
        ("Critic_Score", "Numeric", "Metacritic critical rating (0-100)"),
        ("User_Score", "Numeric", "Users' average rating (0-10)"),
    ],
    widths=[4.0, 3.0, 9.0],
)
add_spacer(doc)
add_para(doc, "Target Variable:", bold=True)
add_para(doc, "None. Clustering is an unsupervised task — the model discovers groups by itself.")

add_heading(doc, "5.4 Data Preparation", 2)
add_bullets(doc, [
    ("1. Load CSV — ", "read the Video Games Sales dataset using pandas."),
    ("2. Clean missing data — ", "convert User_Score to numeric (remove 'tbd'), drop rows missing critic/user/sales fields, keep only positive values."),
    ("3. Select features — ", "build the matrix X = [Critic_Score, User_Score]."),
    ("4. Standardize features — ", "apply StandardScaler to give both features zero mean and unit variance (essential for any distance-based algorithm)."),
    ("5. Fit K-Means — ", "KMeans(n_clusters=3, random_state=42, n_init=10) — n_init=10 runs the algorithm ten times with different seeds and keeps the best result."),
    ("6. Inverse-transform centroids — ", "bring the centroid coordinates back to the original feature space for interpretability."),
])

add_heading(doc, "5.5 Model Training", 2)
add_para(doc, "The model uses scikit-learn's KMeans class with the following configuration:")
add_bullets(doc, [
    ("n_clusters=3 — ", "discover three natural segments in the data."),
    ("random_state=42 — ", "reproducible results."),
    ("n_init=10 — ", "ten independent initializations, keep the one with the lowest inertia."),
])
p = doc.add_paragraph()
r = p.add_run(f"Final inertia (WCSS) after training: {app_r['inertia']}")
r.bold = True
r.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

add_heading(doc, "5.6 Cluster Summary", 2)
summary_rows = [
    (f"C{s['cluster']}", s["count"], f"{s['avg_critic']}", f"{s['avg_user']}",
     f"{s['min_critic']} – {s['max_critic']}", f"{s['min_user']} – {s['max_user']}")
    for s in app_r["summary"]
]
add_table(doc,
    ["Cluster", "Count", "Avg Critic", "Avg User", "Critic Range", "User Range"],
    summary_rows,
    widths=[2.2, 2.2, 2.6, 2.6, 3.4, 3.0],
)
add_spacer(doc)

add_heading(doc, "5.7 Final Centroids", 2)
center_rows = [
    (f"C{c['cluster']}", f"{c['critic_score']}", f"{c['user_score']}")
    for c in app_r["centers"]
]
add_table(doc, ["Cluster", "Critic_Score", "User_Score"], center_rows, widths=[3.5, 5.0, 5.0])
add_spacer(doc)

add_heading(doc, "5.8 Visualization — Scatter Plot with Clusters and Centroids", 2)
scatter_png = os.path.join(_STATIC_DIR, "kmeans_scatter.png")
if os.path.exists(scatter_png):
    doc.add_picture(scatter_png, width=Inches(6.0))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run(f"Figure 2. K-Means clustering of {app_r['dataset_size']} video games. "
                        "Each point is a game colored by cluster; black X markers are the centroids.")
    r.italic = True
    r.font.size = Pt(9.5)

add_heading(doc, "5.9 Interpretation of Results", 2)
add_bullets(doc, [
    ("Cluster 1 — Acclaimed titles — ",
     f"the largest segment ({app_r['summary'][0]['count']} games), high average "
     f"Critic_Score ({app_r['summary'][0]['avg_critic']}) and User_Score "
     f"({app_r['summary'][0]['avg_user']}). Both critics and players praise these games."),
    ("Cluster 2 — Solid mid-tier — ",
     f"the safe-bet segment ({app_r['summary'][1]['count']} games), with moderate "
     f"scores on both axes (Critic {app_r['summary'][1]['avg_critic']}, User "
     f"{app_r['summary'][1]['avg_user']}). Neither praised nor panned."),
    ("Cluster 3 — Underperformers — ",
     f"the smallest segment ({app_r['summary'][2]['count']} games), with low "
     f"scores from both camps (Critic {app_r['summary'][2]['avg_critic']}, User "
     f"{app_r['summary'][2]['avg_user']}). Candidates for discount bundles."),
    ("Why K = 3 works well — ",
     "the three centroids are well separated along the critic-user diagonal; the clusters are stable across runs and easy to characterize."),
    ("Limitations — ",
     "K-Means assumes roughly spherical clusters of similar size, which is only approximately true here. Games at the boundaries between clusters could plausibly belong to either side."),
])

doc.add_page_break()

# ============================ 6. GIT EVIDENCE ============================
add_heading(doc, "6. Git and Repository Evidence", 1)

add_heading(doc, "6.1 Public Repository URL", 2)
add_para(doc, "https://github.com/Niumaster69/Machinlearning")

add_heading(doc, "6.2 Branch Used", 2)
add_table(doc,
    ["Branch", "Purpose", "Commits"],
    [("activityClouster", "Activity 3 — Unsupervised Learning (K-Means Clustering)", "6")],
    widths=[4.5, 7.5, 4.0],
)
add_spacer(doc)

add_heading(doc, "6.3 Evidence of Commits", 2)
add_para(doc, "Progressive commits on branch activityClouster:")
add_table(doc,
    ["Commit", "Description"],
    [
        ("3265ce5", "add Clustering module with manual and sklearn K-Means"),
        ("3550287", "add unsupervised learning basic concepts page"),
        ("d5248bd", "add K-Means manual exercise page with iteration tables"),
        ("817426c", "add clustering application page with scatter plot"),
        ("9ded6ed", "add unsupervised routes and update navbar across all pages"),
        ("293b774", "add Activity 3 report document and generator"),
    ],
    widths=[3.5, 12.5],
)
add_spacer(doc)

add_heading(doc, "6.4 Development Process", 2)
add_para(doc,
    "The development followed a structured, incremental approach: the activityClouster branch "
    "was created from master, the Clustering.py module was implemented first (manual K-Means "
    "plus scikit-learn pipeline), then the three new HTML templates were added (basic concepts, "
    "manual exercise, clustering application), the Flask routes were wired in app.py, and "
    "finally the navbar was updated across every existing template. All commits show "
    "progressive development."
)

doc.add_page_break()

# ============================ 7. RENDER DEPLOYMENT ============================
add_heading(doc, "7. Deployment on Render.com", 1)

add_heading(doc, "7.1 Public URL", 2)
add_para(doc, "[ Paste the Render public URL here after deployment. ]", italic=True)

add_heading(doc, "7.2 Deployment Notes", 2)
add_para(doc,
    "The application is deployed as a Flask web service on Render.com. The repository contains "
    "a requirements.txt listing the Python dependencies (Flask, pandas, numpy, scikit-learn, "
    "matplotlib, seaborn) and a Procfile declaring the web process (gunicorn app:app). On "
    "startup, the Clustering module precomputes both the manual K-Means iterations and the "
    "scikit-learn model, and saves the variance and scatter plots into /static."
)

doc.add_page_break()

# ============================ 8. CONCLUSIONS ============================
add_heading(doc, "8. Conclusions", 1)

add_heading(doc, "8.1 Learning Outcomes", 2)
add_bullets(doc, [
    "Gained hands-on experience implementing the K-Means clustering algorithm both manually (step-by-step Euclidean distances, centroid updates, variance computation) and through scikit-learn.",
    "Learned the key differences between supervised and unsupervised learning: the absence of labels, the use of a geometric/structural criterion (WCSS) rather than an error against ground truth, and the need to interpret the clusters after training.",
    "Understood the importance of feature scaling for distance-based algorithms such as K-Means.",
    "Developed skills in visualizing clustering results through scatter plots with centroids and variance-reduction charts.",
    "Consolidated Flask development skills by extending a multi-page application with a new algorithm section that integrates consistently with the existing supervised learning modules.",
])

add_heading(doc, "8.2 Reflections on Unsupervised Learning", 2)
add_para(doc,
    "Clustering opens a very different perspective compared with the supervised algorithms "
    "implemented in previous activities. Without labels, the evaluation is no longer about "
    "comparing predictions to a known truth but about producing clusters that make sense from "
    "a business or domain point of view. This makes interpretability crucial — the final "
    "centroids must tell a story that a non-technical stakeholder can act upon."
)
add_para(doc,
    "The manual simulation made the mathematics concrete: watching WCSS drop from 5044 to 3787 "
    "across three iterations gave a direct intuition for why the algorithm converges and why "
    "its first iterations are the most impactful."
)

add_heading(doc, "8.3 Comparison with Supervised Models", 2)
add_table(doc,
    ["Aspect", "Supervised (Linear / Logistic / KNN)", "Unsupervised (K-Means)"],
    [
        ("Input", "Features + labels", "Features only"),
        ("Output", "Predicted value/class", "Cluster assignment"),
        ("Training signal", "Error against true label", "Within-cluster variance (WCSS)"),
        ("Evaluation", "Accuracy, R², F1, ROC/AUC", "Inertia, silhouette, domain interpretability"),
        ("Typical goal", "Predict", "Discover structure / segment"),
        ("Project use case", "Predict hours, platform, rating", "Segment games by reception profile"),
    ],
    widths=[4.0, 6.0, 6.0],
)
add_spacer(doc)

add_heading(doc, "8.4 Possible Improvements", 2)
add_bullets(doc, [
    ("Elbow method — ", "plot inertia versus K to justify the choice of K empirically."),
    ("Silhouette score — ", "add an internal validation metric to complement WCSS."),
    ("More features — ", "include regional sales and genre (one-hot encoded) and use PCA to visualize clusters."),
    ("DBSCAN or hierarchical clustering — ", "compare with algorithms that do not require K in advance."),
    ("Interactive dashboard — ", "let the user change K and see clusters update in real time."),
    ("Template inheritance — ", "refactor Flask templates to a base layout to reduce navbar duplication."),
])

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
doc.save(_OUTPUT)
print(f"Report saved to: {_OUTPUT}")
print(f"File size: {os.path.getsize(_OUTPUT)} bytes")
